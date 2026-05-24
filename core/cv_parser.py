import subprocess
import tempfile
from pathlib import Path
from typing import Dict, List

from core.text_utils import (
    extract_email,
    extract_known_skills,
    extract_phone,
    normalize_text,
    top_keywords,
)


class CVParser:
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}

    def parse(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        if suffix not in self.SUPPORTED_EXTENSIONS:
            raise ValueError("Unsupported CV format. Please upload PDF, DOCX, or TXT.")

        if suffix == ".txt":
            text = file_path.read_text(encoding="utf-8", errors="ignore")
        elif suffix == ".pdf":
            text = self._parse_pdf(file_path)
        else:
            text = self._parse_docx(file_path)

        text = normalize_text(text)
        if len(text) < 80:
            raise ValueError(
                "لم أستطع استخراج نص كاف من الملف. لو السيرة الذاتية صورة أو PDF ممسوح ضوئياً، ارفع نسخة نصية PDF أو DOCX أو TXT."
            )
        return text

    def summarize(self, text: str) -> Dict:
        skills = extract_known_skills(text)
        keywords = top_keywords(text, limit=20)
        likely_roles = self._infer_roles(text, keywords)
        return {
            "name": self._guess_name(text),
            "email": extract_email(text),
            "phone": extract_phone(text),
            "skills": skills[:20],
            "likely_roles": likely_roles,
            "years_experience_hint": self._experience_hint(text),
            "text_length": len(text),
        }

    def _parse_pdf(self, file_path: Path) -> str:
        errors = []
        try:
            import fitz

            document = fitz.open(file_path)
            try:
                text = "\n".join(page.get_text() for page in document)
            finally:
                document.close()
            if text.strip():
                return text
            errors.append("PyMuPDF extracted no text")
        except Exception as exc:
            errors.append(f"PyMuPDF: {exc}")

        try:
            from pypdf import PdfReader

            reader = PdfReader(str(file_path))
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages)
            if text.strip():
                return text
            errors.append("pypdf extracted no text")
        except Exception as exc:
            errors.append(f"pypdf: {exc}")

        ocr_text = self._parse_pdf_with_windows_ocr(file_path)
        if ocr_text.strip():
            return ocr_text

        raise RuntimeError(
            "لم أستطع قراءة نص من ملف PDF. غالباً الملف ممسوح ضوئياً أو عبارة عن صور. "
            "ارفع نسخة نصية من الـ PDF أو ملف DOCX/TXT، أو تأكد أن OCR متاح على الجهاز."
        )

    def _parse_pdf_with_windows_ocr(self, file_path: Path) -> str:
        try:
            import fitz
        except Exception:
            return ""

        texts: List[str] = []
        try:
            document = fitz.open(file_path)
            try:
                with tempfile.TemporaryDirectory(prefix="cv_ocr_") as temp_dir:
                    for index, page in enumerate(document):
                        image_path = Path(temp_dir) / f"page_{index + 1}.png"
                        pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                        pixmap.save(str(image_path))
                        page_text = self._ocr_image_with_windows(image_path)
                        if page_text.strip():
                            texts.append(page_text)
            finally:
                document.close()
        except Exception:
            return ""
        return "\n".join(texts)

    def _ocr_image_with_windows(self, image_path: Path) -> str:
        script = r"""
$ErrorActionPreference = 'Stop'
[Console]::OutputEncoding = [System.Text.Encoding]::UTF8
$imagePath = $args[0]
Add-Type -AssemblyName System.Runtime.WindowsRuntime
$null = [Windows.Storage.StorageFile, Windows.Storage, ContentType=WindowsRuntime]
$null = [Windows.Storage.FileAccessMode, Windows.Storage, ContentType=WindowsRuntime]
$null = [Windows.Storage.Streams.IRandomAccessStream, Windows.Storage.Streams, ContentType=WindowsRuntime]
$null = [Windows.Graphics.Imaging.BitmapDecoder, Windows.Graphics.Imaging, ContentType=WindowsRuntime]
$null = [Windows.Graphics.Imaging.SoftwareBitmap, Windows.Graphics.Imaging, ContentType=WindowsRuntime]
$null = [Windows.Media.Ocr.OcrEngine, Windows.Foundation, ContentType=WindowsRuntime]
$asTaskGeneric = ([System.WindowsRuntimeSystemExtensions].GetMethods() | Where-Object {
  $_.Name -eq 'AsTask' -and $_.IsGenericMethod -and $_.GetParameters().Count -eq 1 -and $_.GetParameters()[0].ParameterType.Name -eq 'IAsyncOperation`1'
})[0]
function Await-WinRt($operation, [type]$resultType) {
  $asTask = $asTaskGeneric.MakeGenericMethod($resultType)
  $task = $asTask.Invoke($null, @($operation))
  $task.Wait() | Out-Null
  return $task.Result
}
$file = Await-WinRt ([Windows.Storage.StorageFile]::GetFileFromPathAsync($imagePath)) ([Windows.Storage.StorageFile])
$stream = Await-WinRt ($file.OpenAsync([Windows.Storage.FileAccessMode]::Read)) ([Windows.Storage.Streams.IRandomAccessStream])
$decoder = Await-WinRt ([Windows.Graphics.Imaging.BitmapDecoder]::CreateAsync($stream)) ([Windows.Graphics.Imaging.BitmapDecoder])
$bitmap = Await-WinRt ($decoder.GetSoftwareBitmapAsync()) ([Windows.Graphics.Imaging.SoftwareBitmap])
$engine = [Windows.Media.Ocr.OcrEngine]::TryCreateFromUserProfileLanguages()
if ($null -eq $engine) { exit 2 }
$result = Await-WinRt ($engine.RecognizeAsync($bitmap)) ([Windows.Media.Ocr.OcrResult])
$result.Text
"""
        script_path = None
        try:
            with tempfile.NamedTemporaryFile(
                "w",
                suffix=".ps1",
                prefix="cv_windows_ocr_",
                delete=False,
                encoding="utf-8",
            ) as handle:
                handle.write(script)
                script_path = Path(handle.name)

            completed = subprocess.run(
                [
                    "powershell.exe",
                    "-NoProfile",
                    "-NonInteractive",
                    "-ExecutionPolicy",
                    "Bypass",
                    "-File",
                    str(script_path),
                    str(image_path),
                ],
                capture_output=True,
                encoding="utf-8",
                errors="ignore",
                timeout=45,
                check=False,
            )
        except Exception:
            return ""
        finally:
            if script_path:
                try:
                    script_path.unlink(missing_ok=True)
                except Exception:
                    pass

        if completed.returncode != 0:
            return ""
        return completed.stdout.strip()

    def _parse_docx(self, file_path: Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("DOCX parsing requires python-docx. Install project requirements.") from exc

        document = Document(str(file_path))
        paragraphs = [p.text for p in document.paragraphs]
        table_cells: List[str] = []
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_cells.append(cell.text)
        return "\n".join([*paragraphs, *table_cells])

    def _guess_name(self, text: str) -> str | None:
        for line in text.splitlines()[:8]:
            line = normalize_text(line)
            if 2 <= len(line.split()) <= 4 and "@" not in line and not any(ch.isdigit() for ch in line):
                return line
        return None

    def _infer_roles(self, text: str, keywords: List[str]) -> List[str]:
        lowered = text.lower()
        roles = []
        role_rules = [
            ("Data Analyst", ["data analyst", "excel", "tableau", "power bi", "sql"]),
            ("Data Scientist", ["data scientist", "machine learning", "statistics", "pandas"]),
            ("Machine Learning Engineer", ["machine learning engineer", "deep learning", "tensorflow", "pytorch"]),
            ("Data Engineer", ["data engineer", "etl", "spark", "pipeline"]),
            ("Software Engineer", ["software engineer", "backend", "frontend", "api", "react"]),
            ("Cybersecurity Analyst", ["cybersecurity", "security analyst", "linux", "network security"]),
        ]
        for role, hints in role_rules:
            if any(hint in lowered for hint in hints) or any(hint in keywords for hint in hints):
                roles.append(role)
        return roles[:4] or ["General Technology Candidate"]

    def _experience_hint(self, text: str) -> str | None:
        import re

        match = re.search(r"(\d{1,2})\+?\s*(?:years|yrs)\s+(?:of\s+)?experience", text, re.I)
        return match.group(0) if match else None


cv_parser = CVParser()
